

import streamlit as st
import matplotlib.pyplot as plt
import folium
from streamlit_folium import st_folium
import pandas as pd
import io
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configurar la página para usar ancho completo
st.set_page_config(
    page_title="Dashboard Fiscalía General de la Nación",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

with st.sidebar:
    st.markdown("""
    <div style='margin-top:30px;'>
        <a href='https://www.presidencia.gov.co/' target='_blank'><button class='btn' style='width:100%;margin-bottom:12px;'>Presidencia</button></a>
        <a href='https://www.procuraduria.gov.co/' target='_blank'><button class='btn' style='width:100%;margin-bottom:12px;'>Procuraduría</button></a>
        <a href='https://www.mintic.gov.co/' target='_blank'><button class='btn' style='width:100%;margin-bottom:12px;'>MinTIC</button></a>
        <a href='https://www.fiscalia.gov.co/colombia/' target='_blank'><button class='btn' style='width:100%;margin-bottom:12px;'>Fiscalía</button></a>
        <a href='https://pagina-busqueda-ia.streamlit.app/' target='_blank'><button class='btn' style='width:100%;margin-bottom:12px;background-color:#002855;'>ANÁLISIS DE ARCHIVO CSV - ANÁLISIS DEL DELITO</button></a>
    </div>
    """, unsafe_allow_html=True)

# Encabezado institucional con logo Fiscalía
LOGO_FISCALIA = "https://raw.githubusercontent.com/fredorosemos-hash/PROYECTO-IA/main/assets/logo%20fgn.png"
st.markdown("""
<div style='width:100%;background:#e30613;padding:0;margin:0;'>
    <marquee behavior='scroll' direction='left' style='color:#fff;font-size:1.3rem;font-family:Segoe UI,Arial,sans-serif;font-weight:bold;padding:8px 0;'>
        PROYECTO TALENTO TECH REGION 2 FISCALIA GENERAL DE LA NACION
    </marquee>
</div>
""", unsafe_allow_html=True)
try:
    st.markdown(
        f"""
        <div style='display:flex;justify-content:center;align-items:center;'>
            <img src='{LOGO_FISCALIA}' width='220' alt='Fiscalía General de la Nación'/>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.markdown(
        """
        <div style='text-align:center;'>
            <span style='font-weight:bold;font-size:1.7rem;text-transform:uppercase;'>FISCALÍA GENERAL DE LA NACIÓN SECCIONAL MEDELLÍN</span>
        </div>
        """,
        unsafe_allow_html=True
    )
except Exception:
    st.warning("No se pudo cargar el logo institucional. Verifique la URL o la conexión a internet.")

# Datos de ejemplo
ciudades = ['Bogotá', 'Medellín', 'Cali', 'Barranquilla', 'Cartagena']
delitos_reportados = [1200, 950, 800, 600, 400]
delitos = ['Hurto', 'Homicidio', 'Estafa', 'Secuestro']
delitos_cant = [1800, 700, 500, 250]

# Layout de columnas con mejor distribución
col1, col2, col3 = st.columns([1, 1, 1])
st.markdown("""
<style>
body {
    background: #e0e0e0;
}
.main .block-container {
    padding-left: 1rem;
    padding-right: 1rem;
    max-width: none;
}
.card {
    background: #fff;
    border-radius: 14px;
    box-shadow: 0 2px 12px rgba(0,40,85,0.10);
    border: 2px solid #e30613;
    padding: 18px 16px 12px 16px;
    margin-bottom: 28px;
    width: 100%;
}
.title {
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.15rem;
    color: #002855;
    font-weight: bold;
    margin-bottom: 6px;
    text-transform: uppercase;
    text-align: center;
}
.subtitle {
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.2rem;
    color: #e30613;
    margin-bottom: 12px;
}
.btn {
    background-color: #e30613;
    color: #fff;
    border-radius: 8px;
    padding: 10px 24px;
    font-size: 1rem;
    font-family: 'Segoe UI', Arial, sans-serif;
    border: none;
    cursor: pointer;
    margin-top: 10px;
}
}
</style>
""", unsafe_allow_html=True)
with col1:
    # Gráfico de barras: total de delitos por ciudad
    st.markdown("<div class='card' style='text-align:center;'><div class='title'>Total de delitos por ciudad</div>", unsafe_allow_html=True)
    fig_total, ax_total = plt.subplots(figsize=(6,4), facecolor='#fff')
    ax_total.bar(ciudades, delitos_reportados, color='#e30613', edgecolor='#002855', linewidth=1.5)
    ax_total.set_ylabel('Cantidad', fontsize=12, fontweight='bold', fontname='Arial')
    ax_total.set_xlabel('Ciudad', fontsize=12, fontweight='bold', fontname='Arial')
    ax_total.set_title('Total de delitos por ciudad', color='#002855', fontsize=14, fontweight='bold', fontname='Arial', pad=16)
    ax_total.tick_params(colors='#002855', labelsize=11)
    for label in ax_total.get_xticklabels() + ax_total.get_yticklabels():
        label.set_fontsize(11)
        label.set_fontname('Arial')
    fig_total.patch.set_facecolor('#fff')
    ax_total.spines['top'].set_visible(False)
    ax_total.spines['right'].set_visible(False)
    ax_total.spines['left'].set_color('#002855')
    ax_total.spines['bottom'].set_color('#002855')
    # Etiquetas de valores sobre las barras
    for i, v in enumerate(delitos_reportados):
        ax_total.text(i, v + 20, str(v), ha='center', va='bottom', color='#e30613', fontweight='bold', fontsize=11)
    st.pyplot(fig_total, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

with col2:
    # Gráfico de barras por delitos en cada ciudad
    st.markdown("<div class='card' style='text-align:center;'><div class='title'>Delitos por tipo y ciudad</div>", unsafe_allow_html=True)
    # Datos de ejemplo para delitos por ciudad
    delitos_ciudad = {
        'Bogotá': [700, 300, 150, 50],
        'Medellín': [500, 250, 150, 50],
        'Cali': [400, 200, 150, 50],
        'Barranquilla': [300, 200, 80, 20],
        'Cartagena': [200, 100, 80, 20]
    }
    fig3, ax3 = plt.subplots(figsize=(6,4), facecolor='#fff')
    width = 0.18
    x = range(len(ciudades))
    for i, delito in enumerate(delitos):
        valores = [delitos_ciudad[ciudad][i] for ciudad in ciudades]
        ax3.bar([p + i*width for p in x], valores, width=width, label=delito)
    ax3.set_xticks([p + 1.5*width for p in x])
    ax3.set_xticklabels(ciudades, fontsize=11, fontname='Arial')
    ax3.set_ylabel('Cantidad', fontsize=12, fontweight='bold', fontname='Arial')
    ax3.set_xlabel('Ciudad', fontsize=12, fontweight='bold', fontname='Arial')
    ax3.set_title('Delitos por tipo y ciudad', color='#e30613', fontsize=14, fontweight='bold', fontname='Arial', pad=16)
    ax3.legend(title="Delito", fontsize=9)
    fig3.patch.set_facecolor('#fff')
    ax3.spines['top'].set_visible(False)
    ax3.spines['right'].set_visible(False)
    ax3.spines['left'].set_color('#002855')
    ax3.spines['bottom'].set_color('#002855')
    st.pyplot(fig3, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

with col3:
    # Gráfico de pie con distribución de delitos
    st.markdown("<div class='card' style='text-align:center;'><div class='title'>Distribución de delitos</div>", unsafe_allow_html=True)
    fig2, ax2 = plt.subplots(figsize=(5,4), facecolor='#fff')
    wedges, texts, autotexts = ax2.pie(delitos_cant, labels=delitos, autopct='%1.1f%%', colors=['#ff3b3b','#00eaff','#ffb700','#00ffae'], wedgeprops={'edgecolor':'#fff','linewidth':1.2}, startangle=90)
    for text in texts:
        text.set_fontsize(11)
    fig2.patch.set_facecolor('#fff')
    ax2.set_title('Distribución de Delitos', color='#002855', fontsize=14, fontweight='bold', fontname='Arial', pad=16)
    st.pyplot(fig2, use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

# Sección completa para el mapa (ancho completo)
st.markdown("<div class='card'><div class='title'>Mapa de delitos por ciudad</div>", unsafe_allow_html=True)
coords = {
    'Bogotá': [4.7110, -74.0721],
    'Medellín': [6.2442, -75.5812],
    'Cali': [3.4516, -76.5320],
    'Barranquilla': [10.9685, -74.7813],
    'Cartagena': [10.3910, -75.4794]
}
mapa = folium.Map(location=[4.5709, -74.2973], zoom_start=6)
for ciudad, cantidad in zip(ciudades, delitos_reportados):
    if ciudad in coords:
        folium.Marker(
            coords[ciudad], 
            popup=folium.Popup(f"<b>{ciudad}</b><br>{cantidad} delitos reportados", max_width=200),
            tooltip=f"{ciudad}: {cantidad} delitos",
            icon=folium.Icon(color='red', icon='info-sign')
        ).add_to(mapa)
st_folium(mapa, width=700, height=400, returned_objects=["last_clicked"])
st.markdown("</div>", unsafe_allow_html=True)

# Botón para descargar informe
st.markdown("<div class='card'><div class='title'>Generar informe</div>", unsafe_allow_html=True)
if st.button("Descargar informe", key="neon-btn"):
    import random
    # Generar datos ficticios para 5 años, 5 ciudades, 4 delitos
    anios = [2021, 2022, 2023, 2024, 2025]
    ciudades = ['Bogotá', 'Medellín', 'Cali', 'Barranquilla', 'Cartagena']
    delitos = ['Hurto', 'Homicidio', 'Estafa', 'Secuestro']
    datos_anuales = {}
    for anio in anios:
        datos_anuales[anio] = {}
        for ciudad in ciudades:
            datos_anuales[anio][ciudad] = {}
            for delito in delitos:
                # Generar valores aleatorios
                datos_anuales[anio][ciudad][delito] = random.randint(200, 2500)

    # Crear el documento Word
    doc = Document()
    title = doc.add_heading('INFORME DE DELITOS EN COLOMBIA 2021-2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Fiscalía General de la Nación')
    run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Marco legal y contexto
    marco = doc.add_paragraph()
    marco.add_run('Marco Legal y Contexto Institucional').bold = True
    marco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marco_text = (
        "El presente informe analiza la evolución de los delitos en Colombia durante el periodo 2021-2025, apoyado en datos estadísticos simulados y en el marco de la política criminal nacional. Se consideran las principales ciudades y los delitos de mayor impacto social, con base en la normatividad vigente y las directrices institucionales de la Fiscalía General de la Nación."
    )
    marco_paragraph = doc.add_paragraph(marco_text)
    marco_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Evolución anual por ciudad y delito
    doc.add_heading('Evolución Anual de Delitos por Ciudad y Tipo', level=1)
    for anio in anios:
        doc.add_heading(f'Año {anio}', level=2)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = 'Ciudad'
        for i, delito in enumerate(delitos):
            hdr[i+1].text = delito
        hdr[5].text = 'Total'
        for ciudad in ciudades:
            row = table.add_row().cells
            row[0].text = ciudad
            total_ciudad = 0
            for i, delito in enumerate(delitos):
                val = datos_anuales[anio][ciudad][delito]
                row[i+1].text = str(val)
                total_ciudad += val
            row[5].text = str(total_ciudad)
        doc.add_paragraph('')

    # Gráficos comparativos por año
    for anio in anios:
        fig, ax = plt.subplots(figsize=(6,3))
        valores = [sum(datos_anuales[anio][ciudad][delito] for delito in delitos) for ciudad in ciudades]
        ax.bar(ciudades, valores, color='#e30613', edgecolor='#002855', linewidth=1.5)
        ax.set_ylabel('Delitos', fontsize=12, fontweight='bold', fontname='Arial')
        ax.set_xlabel('Ciudad', fontsize=12, fontweight='bold', fontname='Arial')
        ax.set_title(f'Total de delitos por ciudad en {anio}', fontsize=14, fontweight='bold', fontname='Arial', pad=16)
        fig.tight_layout()
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5.5))
        plt.close(fig)
        doc.add_paragraph('')

    # Gráficos comparativos por tipo de delito
    for delito in delitos:
        fig, ax = plt.subplots(figsize=(6,3))
        valores = [sum(datos_anuales[anio][ciudad][delito] for ciudad in ciudades) for anio in anios]
        ax.plot(anios, valores, marker='o', color='#002855', linewidth=2)
        ax.set_ylabel('Casos', fontsize=12, fontweight='bold', fontname='Arial')
        ax.set_xlabel('Año', fontsize=12, fontweight='bold', fontname='Arial')
        ax.set_title(f'Evolución de {delito} en Colombia (2021-2025)', fontsize=14, fontweight='bold', fontname='Arial', pad=16)
        fig.tight_layout()
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5.5))
        plt.close(fig)
        doc.add_paragraph('')

    # Análisis regional y comparativo
    doc.add_heading('Análisis Regional y Comparativo', level=1)
    for ciudad in ciudades:
        doc.add_heading(f'{ciudad}', level=2)
        texto = (
            f"Durante el periodo 2021-2025, {ciudad} presentó variaciones significativas en la incidencia de delitos. El análisis comparativo muestra tendencias diferenciadas en hurto, homicidio, estafa y secuestro, lo que orienta la focalización de estrategias institucionales."
        )
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Políticas públicas y recomendaciones

    doc.add_heading('Análisis Estratégico para la Mitigación de Delitos', level=1)
    analisis_mitigacion = (
        "A partir de los datos estadísticos y el análisis regional, se proponen estrategias específicas para mitigar los delitos más relevantes en Colombia durante el periodo 2021-2025."
        "\n\nHurto:"
        "\n- Fortalecer la vigilancia policial en zonas de alta incidencia y horarios críticos."
        "\n- Implementar campañas de prevención y educación ciudadana sobre medidas de autoprotección."
        "\n- Promover el uso de tecnologías de videovigilancia y sistemas de alerta temprana en espacios públicos y privados."
        "\n- Fomentar la denuncia y el reporte oportuno de casos para mejorar la reacción institucional."
        "\n\nHomicidio:"
        "\n- Intensificar el trabajo de inteligencia y análisis criminal para identificar patrones y focos de violencia."
        "\n- Desarrollar programas de intervención social en comunidades vulnerables y de alto riesgo."
        "\n- Mejorar la articulación entre Fiscalía, Policía y autoridades locales para la investigación y judicialización efectiva."
        "\n- Fortalecer la protección de testigos y víctimas en procesos judiciales."
        "\n\nEstafa:"
        "\n- Impulsar campañas de sensibilización sobre modalidades de estafa y fraude, especialmente en medios digitales."
        "\n- Reforzar la capacidad de investigación en delitos informáticos y financieros."
        "\n- Promover la cooperación con entidades bancarias y tecnológicas para la detección temprana de fraudes."
        "\n- Fomentar la denuncia y el acompañamiento a víctimas de estafa."
        "\n\nSecuestro:"
        "\n- Fortalecer los grupos especializados en investigación y reacción ante secuestros."
        "\n- Mejorar la coordinación interinstitucional y el intercambio de información entre autoridades nacionales y regionales."
        "\n- Desarrollar estrategias de prevención en zonas rurales y urbanas de alto riesgo."
        "\n- Implementar programas de atención integral a víctimas y familiares."
        "\n\nAcciones transversales:"
        "\n- Promover el uso de herramientas analíticas y tecnológicas para la gestión eficiente de la información criminal."
        "\n- Fomentar la capacitación continua de fiscales e investigadores en análisis criminal y contexto social."
        "\n- Desarrollar campañas de prevención y atención a víctimas en las ciudades con mayor incidencia delictiva."
        "\n- Fortalecer la articulación interinstitucional para la prevención y judicialización de delitos prioritarios."
        "\n\nEstas estrategias deben adaptarse a las particularidades de cada región y evolucionar según las tendencias observadas en los datos, garantizando una respuesta integral y efectiva frente a los retos de la criminalidad en Colombia."
    )
    p = doc.add_paragraph(analisis_mitigacion)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Conclusiones extensas
    doc.add_heading('Conclusiones', level=1)
    conclusiones = (
        "El informe evidencia la importancia de la focalización territorial y el análisis diferenciado por tipo de delito. Las tendencias observadas en los datos simulados permiten orientar la toma de decisiones institucionales y el diseño de políticas públicas para la reducción de la criminalidad. La Fiscalía General de la Nación reitera su compromiso con la seguridad ciudadana y la justicia, promoviendo la innovación y el fortalecimiento de capacidades investigativas en todo el país."
        "\n\nLa evolución anual y regional de los delitos, el análisis comparativo y las recomendaciones presentadas constituyen insumos clave para la gestión estratégica y la mejora continua de la respuesta institucional frente a los retos de la criminalidad en Colombia."
    )
    p = doc.add_paragraph(conclusiones)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sección institucional: Cartilla 5
    cartilla_title = doc.add_paragraph()
    run_cartilla = cartilla_title.add_run('CARTILLA 5: HERRAMIENTAS ANALÍTICAS PARA LA INVESTIGACIÓN Y EL EJERCICIO DE LA ACCIÓN PENAL')
    run_cartilla.bold = True
    cartilla_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cartilla_text = (
        "La política de priorización implica una aproximación analítica y rigurosa a la investigación y al ejercicio de la acción penal. Contribuye al planteamiento de hipótesis delictivas que puedan ser confirmadas o rechazadas a medida que avanza la investigación.\n\n"
        "La resolución 1343 de 2014 establece actividades de priorización respecto a situaciones y casos: fijar un orden en el que serán atendidos; destinar más funcionarios y herramientas; agrupar e investigar casos asociados; aplicar herramientas analíticas en el trámite, investigación y judicialización; enfocar esfuerzos en ciertos casos o hechos delictivos y en algunos presuntos responsables. Tres de las seis actividades proponen el uso del análisis en contexto y la focalización de la investigación y la acción penal en hechos y responsables.\n\n"
        "Esta cartilla presenta cinco herramientas analíticas para analistas criminales, fiscales, investigadores y asistentes de fiscal. Todas son de fácil uso, retoman la forma de analizar la información con la que cuenta la Fiscalía e indican fuentes que pueden contribuir a una investigación más integral. No reemplazan la labor investigativa de la policía judicial.\n\n"
        "Las cinco herramientas parten de tres cambios metodológicos fundamentales:\n"
        "1. Ampliar el foco de la investigación: reconocer que los hechos delictivos no ocurren de manera aislada sino que se explican por su contexto y plantear hipótesis de trabajo para ser confirmadas o rechazadas a través del análisis de la información y evidencia disponibles.\n"
        "2. Disponer de diversas fuentes de información y ser riguroso con su uso: analizar elementos materiales probatorios (EMP), evidencia física (EF) e información de fuentes formales y no formales sobre los hechos y su contexto.\n"
        "3. Utilizar otras disciplinas para explicar fenómenos criminales, situaciones y casos: incluir la aproximación de las ciencias sociales y exactas para comprender el delito e ilustrar la teoría del caso.\n\n"
        "Las herramientas contribuyen a la toma de decisiones de priorización y facilitan el diseño de hipótesis criminales y estrategias de persecución en las diferentes etapas de la investigación, independientemente del régimen procesal. La identificación de fenómenos criminales y la delimitación de situaciones y asociación de casos permiten determinar relaciones entre investigaciones y delitos no denunciados, focalizar la atención en grupos de casos no aislados, distribuir eficazmente recursos y talento humano, y plantear la teoría del caso en etapas preliminares. Las otras tres herramientas aportan elementos para la caracterización de hechos delictivos y actores (víctimas y responsables), útiles en momentos posteriores de la investigación. Permiten focalizar decisiones sobre hechos y personas según criterios de priorización y alimentar la comprensión de la evidencia recolectada.\n\n"
        "Las herramientas son:\n"
        "- Identificación de fenómenos criminales y estudio de resultados.\n"
        "- Delimitación de situación y asociación de casos.\n"
        "- Caracterización de situaciones a partir de prácticas y patrones criminales.\n"
        "- Caracterización de víctimas.\n"
        "- Caracterización de estructuras criminales.\n\n"
        "Estas herramientas permiten investigaciones integrales, rigurosas y estratégicas, orientadas a la priorización y judicialización efectiva."
    )
    cartilla_paragraph = doc.add_paragraph(cartilla_text)
    cartilla_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Guardar y descargar
    buffer = BytesIO()
    doc.save(buffer)
    st.download_button(
        label="Descargar Informe Word",
        data=buffer.getvalue(),
        file_name="Informe_Delitos_Colombia_2021_2025.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown("</div>", unsafe_allow_html=True)
